#!/usr/bin/env python3
"""Build the separate Emerald title page (ECHO_title_page.docx).

Emerald submissions keep author-identifying details on a title page separate from
the anonymised main manuscript. This file carries the title, author, affiliation,
corresponding-author contact, and the funding/conflict statements.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)

# --- Title -----------------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Watching hype arrive: anytime-valid detection of attention "
              "cascades around technology-mediated events on YouTube")
r.bold = True
r.font.size = Pt(16)

doc.add_paragraph()

# --- Author ----------------------------------------------------------------
a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run("Quang-Vinh Dang")
r.bold = True
r.font.size = Pt(13)

aff = doc.add_paragraph()
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = aff.add_run("British University Vietnam, Hung Yen, Vietnam")
r.italic = True
r.font.size = Pt(12)

em = doc.add_paragraph()
em.alignment = WD_ALIGN_PARAGRAPH.CENTER
em.add_run("vinh.dq4@buv.edu.vn").font.size = Pt(12)

doc.add_paragraph()


def block(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# --- Corresponding author + statements -------------------------------------
block("Corresponding author",
      "Quang-Vinh Dang (vinh.dq4@buv.edu.vn), British University Vietnam, "
      "Hung Yen, Vietnam.")
block("Author biography",
      "Quang-Vinh Dang is a researcher at British University Vietnam. His work "
      "spans data science, machine learning and the computational study of "
      "socio-technical systems, with a focus on trustworthy inference from "
      "large-scale behavioural data.")
block("Funding",
      "This research received no specific grant from any funding agency in the "
      "public, commercial or not-for-profit sectors.")
block("Declaration of conflicting interests",
      "The author declares no competing interests.")
block("Data availability",
      "The analysis code and the reproducible collection and analysis pipeline "
      "are openly available. Consistent with the platform's terms and research "
      "ethics, only anonymised, aggregate network and time-series statistics are "
      "reported; commenter identities are irreversibly anonymised at collection "
      "and no verbatim comments are released.")

out = ROOT / "manuscript" / "ECHO_title_page.docx"
doc.save(str(out))
print("wrote", out)
