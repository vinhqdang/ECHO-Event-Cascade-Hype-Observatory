#!/usr/bin/env python3
"""Build the separate Emerald tables file (ECHO_tables.docx), Roman-numeralled."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
RES = json.load(open(ROOT / "results" / "echo_results.json"))

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(11)


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    p.paragraph_format.space_before = Pt(10)


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


def f(x, d=2):
    try:
        return f"{x:.{d}f}"
    except (TypeError, ValueError):
        return str(x)


# ---- Table I: dataset -----------------------------------------------------
title("Table I. Corpus overview")
table(["Event", "Videos", "Comments", "Distinct commenters", "Replies",
       "Innovation-framed videos"],
      [["The Odyssey", 70, "71,020", "53,505", "16,646", 34],
       ["World Cup 2026", 70, "42,867", "36,900", "5,740", 0]])
note("Note: comments sliced to each event's six-week analysis window; commenters "
     "anonymised at collection. Innovation-framed = title/description foregrounds "
     "IMAX/70mm/large-format narrative.")

# ---- Table II: structure --------------------------------------------------
title("Table II. Co-commenter network structural profiles (RQ1)")
p = RES["rq1_structure"]["profiles"]
metrics = [("Nodes", "n_nodes", 0), ("Edges", "n_edges", 0),
           ("Density", "density", 4), ("Modularity", "modularity", 3),
           ("Communities", "n_communities", 0),
           ("Degree assortativity", "assortativity", 3),
           ("Avg. clustering", "avg_clustering", 3),
           ("Core-periphery coeff.", "core_periphery", 3),
           ("Bridge prevalence", "bridge_prevalence", 3),
           ("Largest component frac.", "largest_cc_frac", 3)]
rows = []
for label, key, d in metrics:
    rows.append([label, f(p["odyssey"][key], d), f(p["worldcup"][key], d)])
table(["Metric", "The Odyssey", "World Cup 2026"], rows)
note(f"Graph-level embedding distance = {f(RES['rq1_structure']['embedding_distance'],3)}; "
     f"Laplacian spectral distance = {f(RES['rq1_structure']['spectral_distance'],3)}. "
     "Edge = >=2 shared videos; nodes with degree <2 removed.")

# ---- Table III: detection MC ----------------------------------------------
title("Table III. Sequential detection: Monte-Carlo error control and power (RQ2)")
mc = RES["rq2_detection"]["monte_carlo"]["scenarios"]
labels = {"mixture_poisson": "e-process (Poisson)",
          "mixture_nb": "e-process (NB-robust)",
          "shiryaev_roberts": "e-detector (SR)",
          "fixed_window": "fixed-window (naive)"}
scen_labels = {"null_poisson": "Null (Poisson)",
               "null_overdispersed": "Null (overdispersed)",
               "alt_1p5x": "Shift ×1.5",
               "alt_2p5x": "Shift ×2.5",
               "alt_overdispersed_2x": "Shift ×2 (overdispersed)"}
rows = []
for scen, slab in scen_labels.items():
    for m, mlab in labels.items():
        o = mc[scen][m]
        fa = f(o["alarm_rate"], 3)
        pw = "—" if o["detection_power"] != o["detection_power"] else f(o["detection_power"], 3)
        rows.append([slab, mlab, fa, pw])
table(["Scenario", "Method", "False-alarm prob.", "Detection power"], rows)
note("2,000 replications; alpha = 0.05; baseline rate calibrated to observed data "
     "(63.7 comments/day). For null scenarios 'False-alarm prob.' is the probability "
     "of any alarm; for shift scenarios it is the probability of a pre-change alarm "
     "and 'Detection power' the probability of a valid post-change alarm.")

# ---- Table IV: PSM --------------------------------------------------------
title("Table IV. Causal effect of format-innovation framing on diffusion (RQ3)")
eff = RES["rq3_causal"]["effects"]
rows = []
for o in ("breadth", "volume", "reply_ratio"):
    e = eff[o]
    rows.append([o, f(e["att"], 2), f"[{f(e['ci_low'],2)}, {f(e['ci_high'],2)}]",
                 f(e["naive_diff"], 2), e["n_matched"]])
table(["Outcome", "ATT", "95% CI", "Naive diff.", "Matched pairs"], rows)
note(f"Propensity-score matching on log subscribers, channel video count, upload "
     f"hour, days from anchor, log prior views; 1:1 nearest-neighbour, caliper 0.5 SD. "
     f"{RES['rq3_causal']['n_innovation']} of {RES['rq3_causal']['n_videos']} videos "
     "treated. Max standardised mean difference after matching = "
     f"{f(eff['breadth']['max_smd_after'],3)} (acceptable balance). Reply-chain depth "
     "omitted (structurally capped on YouTube).")

out = ROOT / "manuscript" / "ECHO_tables.docx"
doc.save(str(out))
print("wrote", out)
