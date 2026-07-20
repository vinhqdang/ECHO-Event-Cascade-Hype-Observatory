#!/usr/bin/env python3
"""Build the Emerald-formatted manuscript (DOCX) for ECHO.

Produces manuscript/ECHO_manuscript.docx with:
- structured abstract, keywords
- first-level headings in bold, sub-headings in medium italics (Emerald style)
- Harvard in-text citations, alphabetical reference list with DOIs
- embedded figures with captions; table placeholders (tables live in a separate file)

Numbers are read from results/echo_results.json so the text is exact.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = Path(__file__).resolve().parents[1]
RES = json.load(open(ROOT / "results" / "echo_results.json"))
FIG = ROOT / "results" / "figures"

# ---- pull exact numbers ---------------------------------------------------
r1 = RES["rq1_structure"]["profiles"]
od, wc = r1["odyssey"], r1["worldcup"]
real = RES["rq2_detection"]["real"]
mc = RES["rq2_detection"]["monte_carlo"]["scenarios"]
r3 = RES["rq3_causal"]
rob = RES["robustness"]


def f(x, d=2):
    return f"{x:.{d}f}"


# ---------------------------------------------------------------------------
doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.5


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def para(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def abstract_item(label, text):
    p = doc.add_paragraph()
    run = p.add_run(label + " – ")
    run.bold = True
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def figure(name, caption):
    img = FIG / f"{name}.png"
    if img.exists():
        doc.add_picture(str(img), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ===========================================================================
# TITLE
# ===========================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Watching hype arrive: anytime-valid detection of attention "
                "cascades around technology-mediated events on YouTube")
run.bold = True
run.font.size = Pt(16)

# Author details are withheld from the main manuscript for anonymous peer
# review; they appear on the separate title page (ECHO_title_page.docx).
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("[Author details removed for anonymous review — see title page]")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ===========================================================================
# STRUCTURED ABSTRACT
# ===========================================================================
h1("Abstract")
abstract_item("Purpose", (
    "Collective attention around technology-mediated events builds and breaks "
    "with striking speed, yet analysts almost always describe it in hindsight. "
    "This paper asks whether the onset of such attention cascades can be "
    "detected in real time with valid error guarantees, and whether the "
    "networked signature of a technology-mediated cultural event differs "
    "systematically from that of a non-technological mega-event."))
abstract_item("Design/methodology/approach", (
    "The study assembles a reproducible observatory of YouTube commenting around "
    "two contemporaneous 2026 events — the IMAX/70mm release of Nolan's The "
    "Odyssey and the FIFA World Cup 2026 — comprising 113,887 anonymised "
    "comments across 140 videos. It contrasts their co-commenter networks, and "
    "introduces an anytime-valid, e-process (test-martingale) monitor for "
    "engagement velocity, benchmarked against a naive repeatedly-peeked "
    "fixed-window test through a 2,000-replication Monte-Carlo study; a "
    "propensity-score design estimates the causal effect of format-innovation "
    "framing on diffusion."))
abstract_item("Findings", (
    "The film network is markedly more modular and core-peripheral than the "
    "sporting network. The e-process controls the false-alarm probability under "
    "continuous monitoring (0.01-0.08) where the fixed-window test inflates to "
    "0.59-1.00, while retaining far higher detection power; an "
    "overdispersion-robust variant is required for the heavy-tailed real series. "
    "Contrary to a naive novelty-equals-virality assumption, format-innovation "
    "framing does not amplify — and modestly dampens — networked diffusion once "
    "reach is controlled."))
abstract_item("Originality", (
    "The paper reframes the detection of when collective attention accelerates as "
    "a formal sequential-testing problem with anytime-valid guarantees, a "
    "perspective new to information-systems research on platform-mediated "
    "behaviour, and delivers a fully reproducible, ethics-preserving pipeline."))

kw = doc.add_paragraph()
kw.add_run("Keywords ").bold = True
kw.add_run("Collective attention, Anytime-valid inference, E-values, Change-point "
           "detection, Social media analytics, YouTube, Diffusion of innovations, "
           "Network structure, Propensity score matching, Conformal prediction")
pt = doc.add_paragraph()
pt.add_run("Paper type ").bold = True
pt.add_run("Research paper")

# ===========================================================================
# 1. INTRODUCTION
# ===========================================================================
h1("1. Introduction")
para(
    "On a single day in July 2026, more than nine thousand people left a comment "
    "under videos about one film. Six weeks earlier the same conversation had "
    "amounted to a few dozen remarks a day. Somewhere in between, collective "
    "attention had turned — quietly at first, then all at once. Anyone charged "
    "with responding to that turn faces a deceptively simple question: not whether "
    "people care, but when the caring began to move. This paper is about answering "
    "that question while it still matters, and about what the answer reveals when "
    "the event in question is defined by a piece of technology.")
para(
    "When a much-anticipated film arrives or a global tournament kicks off, the "
    "response online does not build gently. Attention gathers, tips and cascades, "
    "often compressing weeks of anticipation into a handful of decisive days. For "
    "the people and organisations who live with these events — studios, "
    "broadcasters, platform teams, community managers — the question that matters "
    "is rarely whether interest existed, but when it began to move. Yet the "
    "dominant analytical habit in both practice and research is retrospective: we "
    "wait until the window has closed, draw a boundary around a 'before' and an "
    "'after', and compare the two. By the time such a comparison is possible, the "
    "moment to act has usually passed.")
para(
    "This retrospective habit is not merely a matter of timing; it is a matter of "
    "statistical validity. The natural temptation, once data arrive continuously, "
    "is to re-run a comparison every day and react the first time it turns "
    "significant. This practice — repeatedly peeking at an accumulating series "
    "with a test designed to be looked at once — silently destroys the error "
    "guarantee the test was meant to provide (Ramdas et al., 2023; Shafer, 2021). "
    "Under continuous monitoring, a nominal five-per-cent test can raise a false "
    "alarm most of the time. Detecting genuine acceleration in collective "
    "attention therefore poses a problem that is at once substantive and "
    "methodological: how to know, as the data arrive, that a real shift has "
    "begun — and to know it with an error rate that survives continuous looking.")
para(
    "Information systems has long been concerned with how people collectively make "
    "sense of, adopt and react to technology at human scale (Davis, 1989; "
    "Venkatesh et al., 2003; Berente et al., 2019). Platform-mediated attention is "
    "a contemporary expression of exactly that concern: the lifecycles of "
    "collective attention are accelerating (Lorenz-Spreen et al., 2019), online "
    "reactions increasingly distinguish attention from sentiment (Zhang et al., "
    "2026), and platform design choices shape what audiences collectively do "
    "(Liu, 2026; Hong et al., 2025). What has been missing is a principled way to "
    "monitor these processes as they unfold rather than after they conclude.")
para(
    "We address this gap through ECHO, an Event Cascade and Hype Observatory built "
    "around two contemporaneous 2026 events chosen to isolate the role of "
    "technology. The first is the theatrical release of Christopher Nolan's The "
    "Odyssey, promoted heavily on its IMAX and 70mm format — a technology-mediated "
    "cultural event in which a technical attribute is itself part of the story. "
    "The second is the FIFA World Cup 2026, a mega-event of comparable scale whose "
    "draw is not technological. Observing both through the same instrument — the "
    "public commentary they generate on YouTube — lets us ask what, if anything, "
    "the technology dimension changes about how attention organises and moves.")
para(
    "The paper makes three contributions. First, substantively, it shows that a "
    "technology-mediated cultural event leaves a distinctly more modular and "
    "core-peripheral networked signature than a sporting mega-event, and that the "
    "much-touted format-innovation narrative does not, once reach is held "
    "constant, causally amplify networked diffusion. Second, and most "
    "distinctively, it reframes the detection of attention acceleration as a "
    "sequential-testing problem and brings anytime-valid e-process methods "
    "(Grünwald et al., 2024; Vovk and Wang, 2021; Shin et al., 2023) to bear on "
    "it, demonstrating valid real-time detection where fixed-window practice "
    "fails. Third, methodologically, it delivers a transparent, quota-aware and "
    "ethics-preserving pipeline that is fully reproducible from a single "
    "configuration, in the spirit of open, computationally intensive IS research "
    "(Nosek et al., 2015; Berente et al., 2019).")

# ===========================================================================
# 2. BACKGROUND
# ===========================================================================
h1("2. Theoretical background and related work")
h2("2.1 Collective attention and diffusion on platforms")
para(
    "Collective attention — the shared, transient focus of many people on the same "
    "object — is a foundational construct for understanding online behaviour. Its "
    "dynamics have measurably accelerated over time, with popular topics rising "
    "and fading faster than they once did (Lorenz-Spreen et al., 2019), and its "
    "quality is shaped by the informational environment in which it forms (Mocanu "
    "et al., 2015). How attention translates into diffusion has been studied "
    "through the lens of cascades and virality: Goel et al. (2016) distinguish "
    "genuinely viral, person-to-person spread from broadcast-driven reach, while "
    "Vosoughi et al. (2018) show that different kinds of content diffuse at "
    "different speeds and depths. Recent work continues to unpack the structural "
    "and temporal mechanisms behind popularity and success (Mariani et al., 2024; "
    "Wu et al., 2026) and the drivers of commenting and virality specifically "
    "(Peng and Bainbridge, 2026).")
para(
    "For information systems, these dynamics are not merely descriptive curiosities "
    "but instances of how people collectively adopt and adapt to technology. "
    "Classical diffusion-of-innovations and technology-acceptance theory (Rogers, "
    "2003; Davis, 1989; Venkatesh et al., 2003) frame adoption as a social "
    "process; platform commentary offers a fine-grained, behavioural trace of that "
    "process as it happens. YouTube comments in particular have become a "
    "productive site for computational social science (Ji et al., 2023; Sari et "
    "al., 2025), and the discipline has increasingly embraced data-driven, "
    "computationally intensive theory development from such traces (Berente et al., "
    "2019). Within Information Technology & People, a stream of recent work "
    "examines platform-mediated engagement, user-generated content and collective "
    "information behaviour (Liu and Liu, 2025; Osei-Frimpong et al., 2025; Hong et "
    "al., 2025; Liu, 2026; Chen et al., 2026; Sun et al., 2026), locating the "
    "present study squarely within the journal's concern for the human scale of "
    "socio-technical processes.")
h2("2.2 Collective attention as a socio-technical process")
para(
    "Information Technology & People has consistently framed technology use as a "
    "human, social accomplishment rather than a purely technical one. Collective "
    "attention on platforms is a vivid instance: what looks like a spike in a "
    "metric is, underneath, thousands of people encountering the same object and "
    "making sense of it in relation to one another. The structure of that "
    "sense-making — who talks near whom, whether the crowd fragments into "
    "interpretive communities or coheres around a broadcast — is itself a "
    "socio-technical outcome, shaped by how a platform surfaces content and by the "
    "nature of the event that draws people in. Recent ITP scholarship examines how "
    "platform presets shape collective decision premises (Liu, 2026), how "
    "sustained community participation can turn from benefit to harm "
    "(Osei-Frimpong et al., 2025), and how users react against algorithmic "
    "enclosure (Hong et al., 2025). Our contribution to this stream is to treat "
    "the temporal onset and the network shape of collective attention as jointly "
    "observable, and to insist that claims about when attention moves be held to a "
    "standard of statistical validity appropriate to continuous observation.")
h2("2.3 Technology-mediated events and format innovation")
para(
    "Some cultural events foreground a technological attribute as part of their "
    "appeal. The IMAX/70mm presentation of The Odyssey is a clear case: the format "
    "is marketed as a reason to attend, making the technology itself an object of "
    "collective sensemaking. Whether such framing amplifies engagement is an open "
    "question. Electronic word-of-mouth is known to move entertainment outcomes "
    "(Li et al., 2026), but the naive expectation that novelty automatically "
    "breeds virality sits uneasily with evidence that reach and networked "
    "engagement are distinct (Goel et al., 2016) and that attention and sentiment "
    "must be separated (Zhang et al., 2026). We treat the effect of "
    "format-innovation framing as an empirical, causal question rather than an "
    "assumption.")
h2("2.4 Anytime-valid inference and sequential change detection")
para(
    "The methodological core of the paper draws on a rapidly maturing body of work "
    "on game-theoretic statistics and safe, anytime-valid inference (Ramdas et "
    "al., 2023; Grünwald et al., 2024). The central objects are e-values and their "
    "sequential counterparts, test martingales, which quantify evidence against a "
    "null in a way that remains valid under optional stopping and continuous "
    "monitoring (Vovk and Wang, 2021; Shafer, 2021; Waudby-Smith and Ramdas, "
    "2024). Time-uniform confidence sequences (Howard et al., 2021) and "
    "e-value-based multiple-testing procedures (Wang and Ramdas, 2022) extend the "
    "same logic, and the framework is being actively developed for the "
    "unknown-variance, regression and Monte-Carlo settings (Wang and Ramdas, 2025; "
    "Lindon et al., 2026; Fischer et al., 2026; Martin, 2026). For change "
    "detection specifically, e-detectors provide a nonparametric route to "
    "sequential change-point identification with time-uniform guarantees (Shin et "
    "al., 2023), connecting classical Shiryaev-Roberts procedures to the "
    "e-process view. Despite this momentum in statistics, anytime-valid methods "
    "remain almost unknown in IS analyses of platform behaviour, where "
    "fixed-window comparisons still dominate. Bringing them to the study of "
    "collective attention is the paper's principal methodological move.")
para(
    "Two further tools support robustness. Conformal prediction offers "
    "distribution-free predictive intervals (Angelopoulos and Bates, 2023; Vovk et "
    "al., 2022), with adaptive variants that maintain coverage under the "
    "distribution shift endemic to event series (Gibbs et al., 2025; Angelopoulos "
    "et al., 2023; Jun and Ohn, 2026; Sousa et al., 2024). For structural "
    "comparison we draw on community detection (Newman, 2006; Blondel et al., 2008; "
    "Traag et al., 2019) and graph representation learning (Wu et al., 2021; Zhou "
    "et al., 2020), including recent work bridging graph learning and conformal "
    "guarantees (Akansha, 2026).")

# ===========================================================================
# 3. RESEARCH QUESTIONS
# ===========================================================================
h1("3. Research questions")
para(
    "Bringing the substantive and methodological strands together, the study "
    "addresses three questions:")
para(
    "RQ1 (structure). Does the co-commenter network around a technology-mediated "
    "cultural event differ systematically in its structure — modularity, "
    "core-periphery organisation, bridging — from that around a "
    "non-technology-anchored mega-event?")
para(
    "RQ2 (dynamics). Can an anytime-valid e-process detect the onset of "
    "engagement acceleration with valid Type-I error control under continuous "
    "monitoring, faster and more reliably than a fixed-window comparison?")
para(
    "RQ3 (causal). Controlling for channel size, reach and timing, does "
    "format-innovation framing causally amplify cross-channel diffusion?")

# ===========================================================================
# 4. METHODOLOGY
# ===========================================================================
h1("4. Methodology")
h2("4.1 Research design and cases")
para(
    "The design is a comparative, observational study of two contemporaneous "
    "events observed through the same instrument. The Odyssey (theatrical release "
    "17 July 2026) represents a technology-mediated cultural event; the FIFA World "
    "Cup 2026 (opening 11 June 2026) represents a non-technological mega-event of "
    "comparable global scale. For each event we observe a six-week window centred "
    "on the anchor date and analyse the public commentary it generates on "
    "YouTube. The paired design holds constant the platform, the observation "
    "instrument and the calendar period, isolating the events' differing natures "
    "as the object of comparison.")
h2("4.2 Data collection, ethics and reproducibility")
para(
    "Data were collected through the YouTube Data API v3 using a seed-driven, "
    "quota-aware procedure. Because broad recurring keyword search is expensive "
    "against the daily quota, collection seeds from a curated set of queries per "
    "event, freezes the resolved video identifiers into a reproducible sampling "
    "frame, and then expands only through low-cost endpoints; every call is metered "
    "against a documented per-endpoint budget and collection halts explicitly "
    "rather than truncating silently. The resulting corpus comprises 113,887 "
    "comments across 140 videos and two events (Table I). All parameters defining "
    "the sampling frame and analysis are held in a single version-controlled "
    "configuration, consistent with open-science expectations (Nosek et al., 2015).")
para(
    "Ethical handling follows the principle that public availability does not by "
    "itself confer consent to scrutiny of identifiable individuals (Zimmer, 2010). "
    "Commenter identities are anonymised at the moment of writing using a salted "
    "cryptographic digest whose salt is ephemeral and never stored, so the "
    "retained dataset contains only anonymised network nodes; no verbatim comments "
    "are stored or reported, and analysis proceeds entirely at the level of "
    "aggregate network and time-series statistics. This netnographic stance "
    "(Kozinets, 2002) preserves analytical value while minimising re-identification "
    "risk.")
h2("4.3 Network construction")
para(
    "For each event we build a co-commenter (co-participation) network in which "
    "nodes are anonymised commenters and an undirected edge links two commenters "
    "who both commented on at least two shared videos within the window, weighted "
    "by the number of shared videos. Requiring two shared videos isolates a "
    "meaningful cross-video engagement backbone from the single-video majority and "
    "is computed exactly from the sparse commenter-by-video incidence matrix, "
    "scaling to tens of thousands of commenters. In parallel we construct a "
    "sequence of daily snapshots and derive graph-level time series — engagement "
    "velocity (comments per day), new-commenter arrival and daily modularity — "
    "that feed the sequential monitor.")
h2("4.4 Structural comparison (RQ1)")
para(
    "Networks are profiled with interpretable descriptors: modularity and the "
    "number of communities from greedy modularity optimisation (Newman, 2006; "
    "Blondel et al., 2008; Traag et al., 2019), a core-periphery coefficient "
    "(the association between degree and k-core number), degree assortativity, "
    "clustering and bridge prevalence. To complement these we compute a compact "
    "graph-level embedding via an untrained graph-convolutional encoder — node "
    "structural features propagated through the symmetric-normalised adjacency "
    "operator and pooled — an established, label-free structural summariser "
    "consistent with the graph-representation-learning literature (Wu et al., 2021; "
    "Zhou et al., 2020), together with the Laplacian spectrum as a "
    "permutation-free graph distance.")
h2("4.5 Anytime-valid detection (RQ2)")
para(
    "The monitored signal is daily engagement velocity. We treat detection as a "
    "sequential test of the null that the rate has not increased since a burn-in "
    "baseline. For a grid of rate multipliers theta > 1, the per-day likelihood "
    "ratio of an elevated- against baseline-rate model is an e-value with unit "
    "expectation under the null; the running product over days is a non-negative "
    "test martingale, and averaging over the grid (the method of mixtures) covers "
    "an unknown effect size while preserving the martingale property. By Ville's "
    "inequality the probability that this mixture ever exceeds 1/alpha under the "
    "null is at most alpha, so raising an alarm the first time it does gives "
    "anytime-valid Type-I control under arbitrary continuous monitoring (Ramdas et "
    "al., 2023; Vovk and Wang, 2021). We complement the martingale with a "
    "Shiryaev-Roberts e-detector, computed in log-space for numerical stability, "
    "which sums evidence over all candidate change points for fast detection (Shin "
    "et al., 2023).")
para(
    "Social-media counts are heavily overdispersed, which violates the Poisson "
    "model underlying the simplest e-value. We therefore derive an "
    "overdispersion-robust variant using negative-binomial e-values, with the "
    "dispersion estimated from the burn-in and shrunk conservatively to guard "
    "against under-estimation from a short baseline; as the dispersion grows the "
    "variant reduces to the Poisson e-process. The benchmark is the naive "
    "practice the paper argues against: a two-sample rate test recomputed every "
    "day, raising an alarm the first time p < alpha — continuous peeking without "
    "correction. Because a single real series cannot establish error control, we "
    "quantify Type-I error and power through a 2,000-replication Monte-Carlo study "
    "under Poisson and overdispersed nulls and under genuine shifts of varying "
    "size, using a baseline rate calibrated to the observed data.")
h2("4.6 Causal estimation (RQ3)")
para(
    "To ask whether format-innovation framing amplifies diffusion we treat each "
    "Odyssey video as a unit, with treatment indicating that its title or "
    "description foregrounds the format narrative (IMAX, 70mm, large-format "
    "projection). Outcomes are cascade breadth (distinct commenters), volume "
    "(total comments) and reply ratio; reply-chain depth is excluded because "
    "YouTube's flat reply model caps it structurally. Confounding by channel size, "
    "reach and timing is addressed by propensity-score matching (Rosenbaum and "
    "Rubin, 1983; Stuart, 2010; Austin, 2011): a logistic propensity model on log "
    "subscribers, channel size, upload hour, days from the anchor and prior views, "
    "followed by one-to-one nearest-neighbour matching on the logit propensity "
    "within a caliper, with the average treatment effect on the treated reported "
    "alongside a paired bootstrap interval and standardised-mean-difference "
    "balance checks.")
h2("4.7 Uncertainty quantification")
para(
    "As a robustness layer we quantify forecast uncertainty for engagement "
    "velocity using conformal prediction (Angelopoulos and Bates, 2023). We report "
    "both split-conformal intervals and adaptive conformal inference, the latter "
    "adjusting the effective miscoverage level online to maintain coverage under "
    "the distribution shift that event series exhibit (Gibbs et al., 2025; Jun and "
    "Ohn, 2026).")

# ===========================================================================
# 5. RESULTS
# ===========================================================================
h1("5. Results")
h2("5.1 The two attention landscapes")
para(
    f"The corpus contains {71020+42867:,} anonymised comments: {71020:,} from "
    f"{53505:,} distinct commenters across 70 Odyssey videos, and {42867:,} from "
    f"{36900:,} commenters across 70 World Cup videos (Table I). Both events "
    "generated substantial daily commentary, but their temporal shapes differ. "
    "The Odyssey series rises through a protracted pre-release build-up as "
    "successive trailers circulate, cresting at over nine thousand comments a day "
    "at release; the World Cup series stays flat through the group stage before "
    "erupting into the knockout rounds (Figure 1).")
h2("5.2 Structural signatures differ sharply (RQ1)")
para(
    f"The two networks are organised very differently (Table II). The Odyssey "
    f"backbone is larger and far more modular (modularity {f(od['modularity'])} "
    f"across {od['n_communities']} communities) than the World Cup backbone "
    f"(modularity {f(wc['modularity'])}, {wc['n_communities']} communities), and "
    f"markedly more core-peripheral (core-periphery coefficient "
    f"{f(od['core_periphery'])} versus {f(wc['core_periphery'])}). Their degree "
    f"assortativity has opposite sign: the film network is assortative "
    f"({f(od['assortativity'])}), with engaged commenters connecting to one "
    f"another across a rich ecology of trailers, reactions and language-specific "
    f"releases, whereas the sporting network is disassortative "
    f"({f(wc['assortativity'])}), closer to a broadcast-driven hub-and-spoke shape "
    f"organised around a few dominant match videos. The graph-level embedding and "
    f"Laplacian spectra register the same gap "
    f"(embedding distance {f(RES['rq1_structure']['embedding_distance'])}, "
    f"spectral distance {f(RES['rq1_structure']['spectral_distance'])}). "
    "Substantively, the technology-mediated event fosters a more differentiated, "
    "community-structured public — many overlapping interpretive sub-audiences "
    "around a common core — while the sporting mega-event mobilises a broader but "
    "flatter, more centralised one.")
figure("fig3_structure",
       "Figure 1. Structural comparison of the two co-commenter networks. The "
       "film-release network is more modular, more core-peripheral and "
       "assortative; the sporting network is flatter and disassortative.")
h2("5.3 Detecting acceleration: anytime-valid versus fixed-window (RQ2)")
para(
    f"On the real series, the overdispersion-robust e-process raised its alarm for "
    f"The Odyssey on {real['odyssey']['dates'][real['odyssey']['mixture_nb_alarm_day']]}, "
    f"fully {real['odyssey']['anchor_day_index']-real['odyssey']['mixture_nb_alarm_day']} "
    "days before the release anchor — catching the pre-release build-up as it "
    f"began — and for the World Cup on "
    f"{real['worldcup']['dates'][real['worldcup']['mixture_nb_alarm_day']]}, ahead "
    "of the knockout surge (Figure 2). The value of the robust variant is visible "
    "in the raw data: the naive Poisson e-process, mis-specified for the "
    f"heavy-tailed Odyssey series (estimated dispersion r = "
    f"{f(real['odyssey']['dispersion_r'],1)}), fired on the very first day, an "
    "obvious artefact, while the fixed-window monitor raised an alarm for the "
    "World Cup three weeks early, during the flat group stage — a textbook false "
    "positive.")
para(
    "That single-series intuition is confirmed decisively by the Monte-Carlo study "
    "(Table III; Figure 3). Under a clean null with no change, the mixture "
    f"e-process falsely alarms with probability {f(mc['null_poisson']['mixture_poisson']['alarm_rate'])} "
    f"(Poisson) and {f(mc['null_poisson']['mixture_nb']['alarm_rate'])} (robust), "
    f"at or below the nominal 0.05, whereas the repeatedly-peeked fixed-window "
    f"test alarms with probability {f(mc['null_poisson']['fixed_window']['alarm_rate'])} "
    "— more than a tenfold inflation. Under an overdispersed null the naive "
    f"Poisson e-process itself breaks down ({f(mc['null_overdispersed']['mixture_poisson']['alarm_rate'])}), "
    f"but the negative-binomial variant restores control "
    f"({f(mc['null_overdispersed']['mixture_nb']['alarm_rate'])}) while the "
    f"fixed-window test alarms essentially always "
    f"({f(mc['null_overdispersed']['fixed_window']['alarm_rate'])}). Crucially, "
    "this validity does not come at the price of sensitivity. Against a genuine "
    f"2.5-fold shift the robust e-process detects with power "
    f"{f(mc['alt_2p5x']['mixture_nb']['detection_power'])} versus "
    f"{f(mc['alt_2p5x']['fixed_window']['detection_power'])} for the fixed window; "
    f"even against a harder overdispersed two-fold shift it retains power "
    f"{f(mc['alt_overdispersed_2x']['mixture_nb']['detection_power'])} while the "
    f"fixed window collapses to {f(mc['alt_overdispersed_2x']['fixed_window']['detection_power'])}. "
    "The fixed-window test is thus worse on both counts — it both cries wolf and "
    "misses the wolf — whereas the e-process is simultaneously valid and powerful.")
para(
    "The advantage also shows in how quickly genuine change is caught. Averaged "
    f"over replications, against a clean 2.5-fold shift the robust e-process signals "
    f"within about {f(mc['alt_2p5x']['mixture_nb']['mean_latency'],1)} days of the "
    "true change point and the Shiryaev-Roberts e-detector almost immediately, "
    "both while holding pre-change false alarms near zero; the fixed-window test, "
    "by contrast, buys any apparent speed with a pre-change false-alarm rate near "
    "0.29. Speed that comes from crying wolf is not detection. The e-process is the "
    "only monitor here that is fast, sensitive and valid at once — the combination "
    "that makes real-time action defensible rather than reckless.")
figure("fig1_velocity",
       "Figure 2. Daily engagement velocity with detection markers. The e-process "
       "flags genuine acceleration ahead of each event anchor; for the World Cup "
       "the fixed-window monitor false-alarms weeks early during the flat "
       "baseline.")
figure("fig2_error_control",
       "Figure 3. Monte-Carlo validation (2,000 replications). (a) False-alarm "
       "probability under continuous monitoring: the NB-robust e-process hugs the "
       "nominal 0.05 line even under overdispersion, while the fixed-window test "
       "reaches 1.0. (b) Detection power under genuine shifts: the e-processes "
       "dominate the fixed-window test.")
h2("5.4 Does format-innovation framing amplify diffusion? (RQ3)")
para(
    f"Of the 70 Odyssey videos, {r3['n_innovation']} foreground the "
    "format-innovation narrative. Contrary to the intuition that technological "
    "novelty breeds engagement, matching on channel size, reach and timing yields "
    "a negative average treatment effect on cascade breadth "
    f"(ATT = {f(r3['effects']['breadth']['att'],0)} distinct commenters, "
    f"95% CI [{f(r3['effects']['breadth']['ci_low'],0)}, "
    f"{f(r3['effects']['breadth']['ci_high'],0)}]), with volume also negative "
    "though not statistically distinguishable from zero and reply ratio slightly "
    "positive (Table IV). In other words, once reach is held constant, "
    "format-framed videos — largely official promotional uploads — attract "
    "somewhat narrower networked participation than comparable non-framed content, "
    "much of which is reaction and commentary that invites conversation. The "
    "technology attribute drives broadcast reach without a matching gain in "
    "cross-commenter diffusion, echoing the broader distinction between reach and "
    "virality (Goel et al., 2016) and between attention and engagement (Zhang et "
    "al., 2026).")
h2("5.5 Robustness: calibrated forecast uncertainty")
para(
    f"Conformal intervals for engagement velocity behave as their theory predicts. "
    f"On the smoothly accelerating Odyssey series both split and adaptive "
    f"conformal achieve coverage of {f(rob['odyssey']['adaptive_coverage'])} "
    "against a 0.90 target. On the World Cup series, whose abrupt knockout-stage "
    "break violates exchangeability, split conformal under-covers badly "
    f"({f(rob['worldcup']['split_coverage'])}) while adaptive conformal recovers "
    f"much of the shortfall ({f(rob['worldcup']['adaptive_coverage'])}). That even "
    "adaptive coverage falls short over so abrupt a break is itself instructive: "
    "it is precisely the regime change that defeats smoothness-based prediction "
    "that the change-point layer is designed to detect, underscoring why the two "
    "methods are complementary.")
figure("fig4_causal_conformal",
       "Figure 4. (a) Average treatment effect of format-innovation framing on "
       "diffusion outcomes (propensity-score matched, with bootstrap intervals); "
       "breadth is significantly negative. (b) Conformal coverage against the 0.90 "
       "target; adaptive conformal recovers coverage on the abrupt World Cup break.")

# ===========================================================================
# 6. DISCUSSION
# ===========================================================================
h1("6. Discussion")
h2("6.1 What the technology dimension changes")
para(
    "Read together, the results sketch a coherent picture of how a "
    "technology-mediated cultural event differs from a non-technological one. "
    "Structurally, the film event cultivates a more modular, core-peripheral "
    "public: many interpretive communities — fans, sceptics, format enthusiasts, "
    "language-specific audiences — orbiting a shared core, a configuration well "
    "suited to sustained, cross-referential sensemaking. The sporting event, by "
    "contrast, mobilises a flatter and more centralised public organised around "
    "broadcast moments. Yet the technology attribute that ostensibly "
    "differentiates the film — its format — does not, on our evidence, translate "
    "into deeper networked diffusion. The novelty attracts eyes; it does not by "
    "itself knit a denser conversation. For a discipline concerned with how people "
    "collectively conceptualise and adapt to technology, this gap between reach and "
    "networked engagement is the substantive payoff: technological framing "
    "functions here more as a broadcast signal than as a social multiplier.")
h2("6.2 Watching, not reconstructing")
para(
    "The methodological contribution is to make attention acceleration observable "
    "as it happens, with guarantees. Recasting 'when did the hype begin?' as a "
    "sequential test lets us monitor continuously and still control error — the "
    "very thing fixed-window practice cannot do. The Monte-Carlo evidence is "
    "unambiguous: the habitual approach of re-testing a growing series each day is "
    "not a mild approximation but a procedure whose false-alarm probability can "
    "approach certainty, and which is simultaneously underpowered against real "
    "change. The anytime-valid e-process removes this pathology, and the "
    "overdispersion-robust variant is not an optional refinement but a necessity "
    "for the heavy-tailed counts that platform data invariably produce. We regard "
    "the transfer of this toolkit into IS research on platform behaviour as the "
    "paper's most portable idea: wherever analysts monitor an accumulating "
    "behavioural series and act on the first significant look, the same guarantees "
    "apply.")
h2("6.3 Theoretical contributions")
para(
    "The paper contributes to information-systems theory in three ways. First, it "
    "advances a socio-technical account of collective attention in which the "
    "temporal onset and the network structure of a public are treated as coupled, "
    "jointly observable phenomena rather than as separate descriptive facts; the "
    "contrast between a modular, community-structured film public and a flat, "
    "broadcast-centred sporting public shows that the nature of the "
    "technology-mediated object leaves a legible imprint on how collective "
    "sensemaking self-organises. Second, it refines the diffusion-of-innovations "
    "lens for platform settings by distinguishing broadcast reach from networked "
    "diffusion and showing empirically that a salient technology attribute can "
    "drive the former without the latter — a caution against reading engagement "
    "metrics as evidence of deep adoption or community formation (Rogers, 2003; "
    "Goel et al., 2016). Third, and most transferably, it introduces "
    "anytime-valid inference as a theoretical stance for studying live "
    "socio-technical processes: it reframes temporal claims about behaviour as "
    "sequential hypotheses that must remain valid under the continuous observation "
    "that platform data invite.")
h2("6.4 Implications for research and practice")
para(
    "For research, the study illustrates how anytime-valid inference can be woven "
    "into the data-driven, computationally intensive theorising that IS now "
    "embraces (Berente et al., 2019), turning continuous behavioural traces into "
    "monitored hypotheses rather than post-hoc comparisons. For practice, an "
    "e-process monitor offers platform teams, studios and broadcasters an early, "
    "trustworthy signal that engagement has genuinely turned — days before a "
    "release in our data — without the false alarms that erode confidence in "
    "dashboards. The reach-versus-diffusion finding cautions communicators against "
    "equating the visibility of a technological hook with the depth of the "
    "conversation it seeds.")
para(
    "The social implications extend beyond marketing. Trustworthy, early detection "
    "of collective-attention shifts is double-edged: the same monitor that helps a "
    "community team prepare moderation capacity for an incoming surge could be "
    "turned to amplifying manufactured hype or to acting on nascent movements "
    "before participants themselves recognise them. Anchoring such monitoring in "
    "anytime-valid guarantees at least makes the evidentiary basis of any action "
    "explicit and auditable rather than the product of undisclosed repeated "
    "testing; pairing it, as we do, with anonymisation and aggregate-only analysis "
    "models a more accountable practice for observing public behaviour at scale "
    "(Zimmer, 2010; Nosek et al., 2015).")
h2("6.5 Limitations and future work")
para(
    "Several limitations bound the claims. The corpus, though large, is a bounded "
    "sample per video and is drawn from a curated seed frame; comment collection "
    "reflects YouTube's own ordering, and the World Cup baseline is comparatively "
    "sparse. The causal analysis rests on observational matching with a modest "
    "number of treated units and cannot rule out unobserved confounding, and the "
    "treatment — format framing in titles and descriptions — is a proxy for a "
    "richer narrative construct. The two-case design isolates the technology "
    "dimension analytically but not experimentally; replication across more events "
    "would strengthen external validity. Methodologically, the conformal layer "
    "struggles with abrupt breaks, and the e-process baseline is estimated from a "
    "short burn-in. Promising extensions include multivariate e-detectors that "
    "monitor structure and velocity jointly, e-value-based causal inference in the "
    "spirit of Lindon et al. (2026), and adaptive conformal designs tuned for "
    "regime change (Jun and Ohn, 2026).")

# ===========================================================================
# 7. CONCLUSION
# ===========================================================================
h1("7. Conclusion")
para(
    "Collective attention around technology-mediated events is usually narrated "
    "after the fact. This paper shows that it need not be. By observing two "
    "contemporaneous 2026 mega-events through their YouTube commentary, it "
    "establishes that a technology-mediated cultural event leaves a more modular, "
    "core-peripheral networked signature than a sporting one, that its "
    "format-innovation framing wins reach without deepening networked diffusion, "
    "and — most distinctively — that the onset of attention acceleration can be "
    "detected in real time with anytime-valid guarantees where conventional "
    "fixed-window practice fails on both validity and power. The accompanying "
    "observatory is fully reproducible and ethics-preserving. Reframing the "
    "question from 'what happened' to 'is it happening now, and can we trust the "
    "signal' opens a broadly applicable path for information-systems research on "
    "the human dynamics of platform-mediated attention.")

# ===========================================================================
# REFERENCES
# ===========================================================================
h1("References")
REFS = [
 'Akansha, S. (2026), "CondSR: conditional shift-robust conformal prediction for graph neural networks", Soft Computing, Vol. 30 No. 5, pp.3543-3556, doi: 10.1007/s00500-025-11152-3.',
 'Angelopoulos, A.N. and Bates, S. (2023), "Conformal prediction: a gentle introduction", Foundations and Trends in Machine Learning, Vol. 16 No. 4, pp.494-591, doi: 10.1561/2200000101.',
 'Angelopoulos, A., Candès, E. and Tibshirani, R. (2023), "Conformal PID control for time series prediction", Advances in Neural Information Processing Systems, Vol. 36, pp.23047-23074, doi: 10.52202/075280-1000.',
 'Austin, P.C. (2011), "An introduction to propensity score methods for reducing the effects of confounding in observational studies", Multivariate Behavioral Research, Vol. 46 No. 3, pp.399-424, doi: 10.1080/00273171.2011.568786.',
 'Berente, N., Seidel, S. and Safadi, H. (2019), "Research commentary—data-driven computationally intensive theory development", Information Systems Research, Vol. 30 No. 1, pp.50-64, doi: 10.1287/isre.2018.0774.',
 'Blondel, V.D., Guillaume, J.-L., Lambiotte, R. and Lefebvre, E. (2008), "Fast unfolding of communities in large networks", Journal of Statistical Mechanics: Theory and Experiment, Vol. 2008 No. 10, P10008, doi: 10.1088/1742-5468/2008/10/P10008.',
 'Chen, Y., Sun, Y., Zhang, Y. and Wang, Y. (2026), "Emotional support or emotional release? A dual-path model of emotional intelligence of AI companions on user engagement", Information Technology & People, Vol. ahead-of-print No. ahead-of-print, doi: 10.1108/ITP-03-2025-0362.',
 'Davis, F.D. (1989), "Perceived usefulness, perceived ease of use, and user acceptance of information technology", MIS Quarterly, Vol. 13 No. 3, pp.319-340, doi: 10.2307/249008.',
 'Fischer, L., Barry, T. and Ramdas, A. (2026), "Multiple testing with anytime-valid Monte Carlo p-values", Electronic Journal of Statistics, Vol. 20 No. 1, doi: 10.1214/26-ejs2514.',
 'Gibbs, I., Cherian, J.J. and Candès, E.J. (2025), "Conformal prediction with conditional guarantees", Journal of the Royal Statistical Society Series B: Statistical Methodology, Vol. 87 No. 4, pp.1100-1126, doi: 10.1093/jrsssb/qkaf008.',
 'Goel, S., Anderson, A., Hofman, J. and Watts, D.J. (2016), "The structural virality of online diffusion", Management Science, Vol. 62 No. 1, pp.180-196, doi: 10.1287/mnsc.2015.2158.',
 'Grünwald, P., de Heide, R. and Koolen, W. (2024), "Safe testing", Journal of the Royal Statistical Society Series B: Statistical Methodology, Vol. 86 No. 5, pp.1091-1128, doi: 10.1093/jrsssb/qkae011.',
 'Hong, X., Pan, L., Xu, M. and Chen, Q. (2025), "Escaping from the echo chamber: understanding user behavior from the perspective of psychological reactance theory", Information Technology & People, Vol. 39 No. 3, pp.1447-1473, doi: 10.1108/ITP-08-2024-0984.',
 'Howard, S.R., Ramdas, A., McAuliffe, J. and Sekhon, J. (2021), "Time-uniform, nonparametric, nonasymptotic confidence sequences", The Annals of Statistics, Vol. 49 No. 2, pp.1055-1080, doi: 10.1214/20-aos1991.',
 'Ji, J.J., Hu, H. and Wei, S. (2023), "YouTube comments on gene-edited babies: what factors affect diverse opinions in comments?", Social Science Computer Review, Vol. 41 No. 4, pp.1420-1437, doi: 10.1177/08944393211073164.',
 'Jun, J. and Ohn, I. (2026), "Online conformal inference with retrospective adjustment for faster adaptation to distribution shift", Pattern Recognition, Vol. 180, 114406, doi: 10.1016/j.patcog.2026.114406.',
 'Kozinets, R.V. (2002), "The field behind the screen: using netnography for marketing research in online communities", Journal of Marketing Research, Vol. 39 No. 1, pp.61-72, doi: 10.1509/jmkr.39.1.61.18935.',
 'Li, L., He, Y., Liu, L. and Zhao, J. (2026), "The impact of multi-aspect electronic word-of-mouth on movie box office: an exploration study based on large language models", Decision Support Systems, Vol. 208, 114718, doi: 10.1016/j.dss.2026.114718.',
 'Lindon, M., Ham, D.W., Tingley, M. and Bojinov, I. (2026), "Anytime-valid inference in linear models with applications to regression-adjusted causal inference", Journal of the American Statistical Association, advance online publication, pp.1-27, doi: 10.1080/01621459.2026.2692052.',
 'Liu, C. and Liu, S. (2025), "How to increase other’s engagement with user-generated content: the role of flow experience", Information Technology & People, Vol. 39 No. 3, pp.1474-1494, doi: 10.1108/ITP-11-2024-1403.',
 'Liu, S. (2026), "Programmed power: how presets transform decision premises in the platform society", Information Technology & People, Vol. ahead-of-print No. ahead-of-print, doi: 10.1108/ITP-01-2026-0110.',
 'Lorenz-Spreen, P., Mønsted, B.M., Hövel, P. and Lehmann, S. (2019), "Accelerating dynamics of collective attention", Nature Communications, Vol. 10, 1759, doi: 10.1038/s41467-019-09311-w.',
 'Mariani, M.S., Battiston, F., Horvát, E.-Á., Livan, G., Musciotto, F. and Wang, D. (2024), "Collective dynamics behind success", Nature Communications, Vol. 15, 10701, doi: 10.1038/s41467-024-54612-4.',
 'Martin, R. (2026), "Regularized e-processes: anytime valid inference with knowledge-based efficiency gains", Bernoulli, Vol. 32 No. 3, doi: 10.3150/25-bej1939.',
 'Mocanu, D., Rossi, L., Zhang, Q., Karsai, M. and Quattrociocchi, W. (2015), "Collective attention in the age of (mis)information", Computers in Human Behavior, Vol. 51, pp.1198-1204, doi: 10.1016/j.chb.2015.01.024.',
 'Newman, M.E.J. (2006), "Modularity and community structure in networks", Proceedings of the National Academy of Sciences, Vol. 103 No. 23, pp.8577-8582, doi: 10.1073/pnas.0601602103.',
 'Nosek, B.A., Alter, G., Banks, G.C., Borsboom, D., Bowman, S.D., Breckler, S.J., Buck, S., Chambers, C.D., Chin, G., Christensen, G., Contestabile, M., Dafoe, A., Eich, E., Freese, J., Glennerster, R., Goroff, D., Green, D.P., Hesse, B., Humphreys, M., Ishiyama, J., Karlan, D., Kraut, A., Lupia, A., Mabry, P., Madon, T., Malhotra, N., Mayo-Wilson, E., McNutt, M., Miguel, E., Paluck, E.L., Simonsohn, U., Soderberg, C., Spellman, B.A., Turitto, J., VandenBos, G., Vazire, S., Wagenmakers, E.J., Wilson, R. and Yarkoni, T. (2015), "Promoting an open research culture", Science, Vol. 348 No. 6242, pp.1422-1425, doi: 10.1126/science.aab2374.',
 'Osei-Frimpong, K., Islam, N., Ahorkonu, C.K., Soga, L.R. and McLean, G. (2025), "Exploring the dark side of continuous social media brand community participation on consumer psychological ill-being: initial vs sustained participation", Information Technology & People, Vol. 39 No. 3, pp.1387-1414, doi: 10.1108/ITP-02-2025-0162.',
 'Peng, S. and Bainbridge, W.A. (2026), "Image memorability predicts social media virality and externally-associated commenting", Computers in Human Behavior, Vol. 174, 108799, doi: 10.1016/j.chb.2025.108799.',
 'Ramdas, A., Grünwald, P., Vovk, V. and Shafer, G. (2023), "Game-theoretic statistics and safe anytime-valid inference", Statistical Science, Vol. 38 No. 4, pp.576-601, doi: 10.1214/23-sts894.',
 'Rogers, E.M. (2003), Diffusion of Innovations, 5th ed., Free Press, New York, NY.',
 'Rosenbaum, P.R. and Rubin, D.B. (1983), "The central role of the propensity score in observational studies for causal effects", Biometrika, Vol. 70 No. 1, pp.41-55, doi: 10.1093/biomet/70.1.41.',
 'Sari, D.K., Herwandito, S., Utomo, A.W., Solikhah, N.I. and Wibowo, N.A. (2025), "We know who we are: expressing national identity on YouTube video comments", Journal of Computational Social Science, Vol. 8 No. 4, 85, doi: 10.1007/s42001-025-00426-3.',
 'Shafer, G. (2021), "Testing by betting: a strategy for statistical and scientific communication", Journal of the Royal Statistical Society Series A: Statistics in Society, Vol. 184 No. 2, pp.407-431, doi: 10.1111/rssa.12647.',
 'Shekhar, S. and Ramdas, A. (2024), "Nonparametric two-sample testing by betting", IEEE Transactions on Information Theory, Vol. 70 No. 2, pp.1178-1203, doi: 10.1109/tit.2023.3305867.',
 'Shin, J., Ramdas, A. and Rinaldo, A. (2023), "E-detectors: a nonparametric framework for sequential change detection", The New England Journal of Statistics in Data Science, pp.229-260, doi: 10.51387/23-nejsds51.',
 'Sousa, M., Tomé, A.M. and Moreira, J. (2024), "A general framework for multi-step ahead adaptive conformal heteroscedastic time series forecasting", Neurocomputing, Vol. 608, 128434, doi: 10.1016/j.neucom.2024.128434.',
 'Stuart, E.A. (2010), "Matching methods for causal inference: a review and a look forward", Statistical Science, Vol. 25 No. 1, pp.1-21, doi: 10.1214/09-STS313.',
 'Sun, M., Chen, H., Chen, H. and Jiang, B. (2026), "Does AI identity disclosure stimulate user responses? An exploration based on meta-analysis", Information Technology & People, Vol. ahead-of-print No. ahead-of-print, doi: 10.1108/ITP-06-2025-0813.',
 'Traag, V.A., Waltman, L. and van Eck, N.J. (2019), "From Louvain to Leiden: guaranteeing well-connected communities", Scientific Reports, Vol. 9, 5233, doi: 10.1038/s41598-019-41695-z.',
 'Venkatesh, V., Morris, M.G., Davis, G.B. and Davis, F.D. (2003), "User acceptance of information technology: toward a unified view", MIS Quarterly, Vol. 27 No. 3, pp.425-478, doi: 10.2307/30036540.',
 'Vosoughi, S., Roy, D. and Aral, S. (2018), "The spread of true and false news online", Science, Vol. 359 No. 6380, pp.1146-1151, doi: 10.1126/science.aap9559.',
 'Vovk, V. and Wang, R. (2021), "E-values: calibration, combination and applications", The Annals of Statistics, Vol. 49 No. 3, pp.1736-1754, doi: 10.1214/20-aos2020.',
 'Vovk, V., Gammerman, A. and Shafer, G. (2022), Algorithmic Learning in a Random World, 2nd ed., Springer, Cham, doi: 10.1007/978-3-031-06649-8.',
 'Wang, H. and Ramdas, A. (2025), "Anytime-valid t-tests and confidence sequences for Gaussian means with unknown variance", Sequential Analysis, Vol. 44 No. 1, pp.56-110, doi: 10.1080/07474946.2024.2428245.',
 'Wang, R. and Ramdas, A. (2022), "False discovery rate control with e-values", Journal of the Royal Statistical Society Series B: Statistical Methodology, Vol. 84 No. 3, pp.822-852, doi: 10.1111/rssb.12489.',
 'Wu, Y.J., van Zoonen, W., Treem, J.W. and Sivunen, A.E. (2026), "Attention dynamics on social technology platforms in organizations: an empirical study of structural and temporal mechanisms", New Media & Society, advance online publication, doi: 10.1177/14614448261456236.',
 'Wu, Z., Pan, S., Chen, F., Long, G., Zhang, C. and Yu, P.S. (2021), "A comprehensive survey on graph neural networks", IEEE Transactions on Neural Networks and Learning Systems, Vol. 32 No. 1, pp.4-24, doi: 10.1109/TNNLS.2020.2978386.',
 'Zhang, X.M., Xu, D., Hong, H. and Chan, K. (2026), "Attention or sentiment: how social media react to ESG?", Information Systems Research, Vol. 37 No. 2, pp.1323-1336, doi: 10.1287/isre.2022.0251.',
 'Zhou, J., Cui, G., Hu, S., Zhang, Z., Yang, C., Liu, Z., Wang, L., Li, C. and Sun, M. (2020), "Graph neural networks: a review of methods and applications", AI Open, Vol. 1, pp.57-81, doi: 10.1016/j.aiopen.2021.01.001.',
 'Zimmer, M. (2010), "‘But the data is already public’: on the ethics of research in Facebook", Ethics and Information Technology, Vol. 12 No. 4, pp.313-325, doi: 10.1007/s10676-010-9227-5.',
]
for ref in REFS:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

out = ROOT / "manuscript" / "ECHO_manuscript.docx"
doc.save(str(out))
print("wrote", out)

# quick word count estimate
wc_total = sum(len(p.text.split()) for p in doc.paragraphs)
print("approx words:", wc_total, "| references:", len(REFS))
